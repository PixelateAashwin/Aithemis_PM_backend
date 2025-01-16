
from fastapi import FastAPI, File, UploadFile, HTTPException
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document  
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from pydantic import BaseModel
import numpy as np
import os
import tempfile
import logging
from docx import Document  # Import the Document class from python-docx

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

# Initialize the FAISS vectorstore
vectorstore = None

class SearchQuery(BaseModel):
    query: str
    parsedData: list 

# Use a more advanced model for embeddings
class CustomEmbeddings:
    def __init__(self, model_name="sentence-transformers/all-MiniLM-L6-v2"):
        self.model = SentenceTransformer(model_name)
    
    def embed_documents(self, texts):
        """Embed a list of documents"""
        if not texts:
            return []
        embeddings = self.model.encode(texts, convert_to_numpy=True)
        return [embedding.tolist() for embedding in embeddings]
    
    def embed_query(self, text):
        """Embed a query"""
        if not isinstance(text, str):
            raise ValueError("Query must be a string")
        return self.model.encode(text, convert_to_numpy=True).tolist()
    
    def embed_search(self, text):
        """Compatibility method for similarity search"""
        return self.embed_query(text)
    
    def __call__(self, text):
        """Make the class callable - required for FAISS"""
        if isinstance(text, list):
            return self.embed_documents(text)
        return self.embed_query(text)

# Function to normalize embeddings
def normalize_embeddings(embeddings):
    """Normalize embeddings to unit vectors"""
    norm = np.linalg.norm(embeddings, axis=1, keepdims=True)
    return embeddings / norm if norm.all() else embeddings

# Update chunking strategy to preserve context
def chunk_document(document_text, chunk_size=1000, overlap_size=200):
    """Chunk the document into semantically aware chunks"""
    chunks = []
    for i in range(0, len(document_text), chunk_size - overlap_size):
        chunk = document_text[i:i + chunk_size]
        chunks.append(chunk)
    return chunks

# Initialize embeddings
try:
    embeddings = CustomEmbeddings()
    # logger.info("Successfully initialized embeddings model")
except Exception as e:
    # logger.error(f"Error initializing embeddings: {e}")
    raise

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    if not (file.filename.endswith('.pdf') or file.filename.endswith('.docx')):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    
    try:
        # Create a temporary file to store the uploaded document
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf' if file.filename.endswith('.pdf') else '.docx') as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_path = temp_file.name

        # Load and process the document
        documents = []
        extracted_text = ""
        if file.filename.endswith('.pdf'):
            loader = PyPDFLoader(temp_path)
            documents = loader.load()
            extracted_text = "\n".join(doc.page_content for doc in documents)
            # logger.info(f"Loaded PDF with {len(documents)} pages")
        elif file.filename.endswith('.docx'):
            extracted_text = load_docx(temp_path)
            doc = Document()
            doc.page_content = extracted_text 
            doc.metadata = {"filename": file.filename}
            documents = [doc]
            # logger.info(f"Loaded DOCX with {len(extracted_text.splitlines())} paragraphs")

        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=100,
            separators=["\n\n", "\n", ".", " "]
        )
        chunks = text_splitter.split_documents(documents)
        # logger.info(f"Split into {len(chunks)} chunks")

        # Generate embeddings for the chunks
        embeddings_list = embeddings.embed_documents([chunk.page_content for chunk in chunks])

        # Create a dictionary with parsed text and embeddings
        parsed_text = extracted_text
        embeddings_data = embeddings_list  

        # Clean up temporary file
        os.unlink(temp_path)

        # Return both parsed text and embeddings
        return {
            "message": "File uploaded and processed successfully",
            "document_count": len(documents),
            "chunk_count": len(chunks),
            "parsed_text": parsed_text,
            "embeddings": embeddings_data 
        }

    except Exception as e:
        logger.error(f"Error processing upload: {e}")
        # Clean up temp file if it exists
        if 'temp_path' in locals():
            os.unlink(temp_path)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/search")
async def search_documents(search_query: SearchQuery):
    try:
        # Embed the query
        query_embedding = embeddings.embed_query(search_query.query)

        # Ensure query_embedding is a 2D array (1, embedding_size)
        query_embedding = np.array(query_embedding).reshape(1, -1)

        # Minimum similarity threshold
        SIMILARITY_THRESHOLD = 0.1

        results = []
        # Iterate over parsedData 
        for doc_data in search_query.parsedData:
            # Extract embeddings array and text content
            document_embeddings = np.array(doc_data['embedding'])
            document_text = doc_data.get('text', '')

            # Ensure document_embeddings is a 2D array
            document_embeddings = document_embeddings.reshape(-1, len(query_embedding[0]))
            document_embeddings = normalize_embeddings(document_embeddings)

            # Store all matches above threshold
            matches = [] 
            chunk_size = 500

            # Split document text into chunks if available
            chunks = []
            if document_text:
                chunks = [document_text[i:i + chunk_size] for i in range(0, len(document_text), chunk_size)]

            # Calculate similarity for each embedding and store if above threshold
            similarities = cosine_similarity(query_embedding, document_embeddings)

            for i, similarity in enumerate(similarities[0]):
                if similarity >= SIMILARITY_THRESHOLD:
                    match = {
                        "chunk_index": i,
                        "page_number": i + 1,  # Assuming 1-based page numbering
                        "similarity": float(similarity),
                        "text": chunks[i] if chunks and i < len(chunks) else f"Chunk {i}",
                        "metadata": doc_data.get('metadata', {}),
                        "document_name": doc_data['metadata'].get('name', 'Unknown Document'),
                    }
                    matches.append(match)

            # Sort matches by similarity
            matches.sort(key=lambda x: x['similarity'], reverse=True)

            # If we found any matches, add them to results
            if matches:
                doc_result = {
                    "document_name": doc_data['metadata'].get('name', 'Unknown Document'),
                    "document_metadata": doc_data.get('metadata', {}),
                    "matches": matches
                }
                results.append(doc_result)

        # Sort documents by their best match similarity
        results.sort(key=lambda x: max(m['similarity'] for m in x['matches']) if x['matches'] else 0, reverse=True)

        return {
            "query": search_query.query,
            "results": results,
            "total_results": len(results),
            "total_matches": sum(len(doc['matches']) for doc in results)
        }
    except Exception as e:
        logger.error(f"Error during search: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health_check():
    return {"status": "healthy", "embeddings_loaded": embeddings is not None}

def load_docx(file_path):
    """Load content from a DOCX file."""
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000) 
